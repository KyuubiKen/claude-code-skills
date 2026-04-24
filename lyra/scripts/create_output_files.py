#!/usr/bin/env python3
"""
Lyra output file creator.
Saves optimized prompts to Desktop as both .md and .docx (raw markdown text).
"""

import argparse
import os
import sys
from pathlib import Path


def create_md_file(content: str, name: str, desktop: Path) -> Path:
    path = desktop / f"{name}.md"
    path.write_text(content, encoding="utf-8")
    return path


def create_docx_file(content: str, name: str, desktop: Path) -> Path:
    path = desktop / f"{name}.docx"

    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Remove default empty paragraph
        for para in doc.paragraphs:
            p = para._element
            p.getparent().remove(p)

        # Write each line as its own paragraph to preserve raw markdown
        for line in content.splitlines():
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(11)

        doc.save(str(path))

    except ImportError:
        # Fallback: write raw markdown text into a plain .docx-named file
        # Not a true .docx, but openable as text
        path.write_text(content, encoding="utf-8")
        print("Warning: python-docx not installed. Saved raw text as .docx.")
        print("To install: pip3 install python-docx")

    return path


def main():
    parser = argparse.ArgumentParser(description="Create Lyra output files on Desktop")
    parser.add_argument("--content", required=True, help="The optimized prompt content")
    parser.add_argument("--name", required=True, help="Base filename (snake_case topic/intent)")
    args = parser.parse_args()

    desktop = Path.home() / "Desktop"
    if not desktop.exists():
        print(f"Error: Desktop not found at {desktop}", file=sys.stderr)
        sys.exit(1)

    md_path = create_md_file(args.content, args.name, desktop)
    docx_path = create_docx_file(args.content, args.name, desktop)

    print(f"Saved: {md_path}")
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    main()
